from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from datetime import datetime


def export_to_word(analyse_tekst: str, vergader_type: str, output_path: str) -> str:
    doc = Document()

    # Title
    title = doc.add_heading(f"Vergadervoorbereiding — {vergader_type}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    datum = doc.add_paragraph(f"Gegenereerd op: {datetime.now().strftime('%d-%m-%Y %H:%M')}")
    datum.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    datum.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # Parse and render markdown-ish content
    lines = analyse_tekst.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("**") and stripped.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.bold = True
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.", stripped):
            doc.add_paragraph(stripped, style="List Number")
        else:
            # Handle inline bold (**text**)
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*.*?\*\*)", stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    doc.save(output_path)
    return output_path
